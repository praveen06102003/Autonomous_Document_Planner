"""
doc_builder.py — turns an approved plan (title + summary + tasks + diagram)
into a polished .docx file using python-docx. This only runs ONCE, after
the user has confirmed the plan is final — not on every planning/refine step.
"""

import os
import base64
import requests
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def _fetch_diagram_image(diagram_text: str, plan_id: str) -> str | None:
    """
    Converts Mermaid diagram text into a PNG using the free mermaid.ink
    rendering service, saves it locally, and returns the filepath.
    Returns None if the diagram is empty or the fetch fails — the doc
    is still generated without a diagram rather than failing entirely.
    """
    if not diagram_text or not diagram_text.strip():
        return None

    try:
        encoded = base64.urlsafe_b64encode(diagram_text.encode("utf-8")).decode("utf-8")
        url = f"https://mermaid.ink/img/{encoded}?type=png"
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            image_path = os.path.join(OUTPUT_DIR, f"{plan_id}_diagram.png")
            with open(image_path, "wb") as f:
                f.write(response.content)
            return image_path
    except Exception as e:
        print("Diagram image fetch failed:", e)

    return None


def build_document(plan_id: str, title: str, summary: str, tasks: list[dict], diagram: str = "") -> str:
    """
    Builds a Word doc styled like a real document: title, summary
    paragraph, an "Action Plan" section with each task as a heading +
    notes + deadline, and an embedded architecture diagram if provided.
    Returns the filename (not full path) so main.py can build a download URL.
    """
    doc = Document()

    # Title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata line
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Generated on {datetime.now().strftime('%d %b %Y, %I:%M %p')}")
    meta_run.italic = True
    meta_run.font.size = Pt(9)

    doc.add_paragraph()  # spacer

    # Summary paragraph (like the chat summary, shown as document intro)
    if summary:
        summary_para = doc.add_paragraph(summary)
        summary_para.runs[0].font.size = Pt(11)
        doc.add_paragraph()  # spacer

    # Action Plan section
    doc.add_heading("Action Plan", level=2)

    for task in tasks:
        # Task title as a sub-heading, with deadline inline
        task_heading = doc.add_paragraph()
        task_run = task_heading.add_run(task.get("task", ""))
        task_run.bold = True
        task_run.font.size = Pt(12)

        deadline = task.get("deadline", "")
        if deadline:
            deadline_run = task_heading.add_run(f"   —   {deadline}")
            deadline_run.italic = True
            deadline_run.font.size = Pt(9)

        # Notes as a regular paragraph underneath
        notes = task.get("notes", "")
        if notes:
            notes_para = doc.add_paragraph(notes)
            notes_para.runs[0].font.size = Pt(10.5)

        doc.add_paragraph()  # spacer between tasks

    # Architecture Diagram section (only if a diagram was generated)
    if diagram:
        image_path = _fetch_diagram_image(diagram, plan_id)
        if image_path:
            doc.add_heading("Architecture Diagram", level=2)
            doc.add_picture(image_path, width=Inches(5.5))

    filename = f"{plan_id}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)

    return filename